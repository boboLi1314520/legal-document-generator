"""
Word 文档生成服务 - 重构版
支持模板变量替换
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, List
import os
import re
import copy
from docx.oxml.ns import qn


class DocxService:
    """Word 文档生成服务"""

    def generate_complaint(self, case: dict, content: str, output_path: str):
        """生成民事起诉状"""
        doc = Document()

        # 标题
        title = doc.add_paragraph("民事起诉状")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(22)
        title_run.font.bold = True

        doc.add_paragraph()

        # 原告信息
        p = doc.add_paragraph()
        p.add_run("原告：").bold = True
        p.add_run(case.get("plaintiff", "待填写"))

        # 被告信息
        p = doc.add_paragraph()
        p.add_run("被告：").bold = True
        p.add_run(case.get("defendant", "待填写"))

        doc.add_paragraph()

        # 诉讼请求
        p = doc.add_paragraph()
        p.add_run("诉讼请求：").bold = True
        doc.add_paragraph("1. 判令被告偿还债务本金及利息")
        doc.add_paragraph("2. 判令被告承担本案诉讼费用")

        doc.add_paragraph()

        # 事实与理由
        p = doc.add_paragraph()
        p.add_run("事实与理由：").bold = True

        facts = case.get("facts", content)
        p = doc.add_paragraph(facts if facts else "待填写")

        doc.add_paragraph()
        doc.add_paragraph()

        # 此致
        p = doc.add_paragraph()
        p.add_run("此致").bold = True

        court = case.get("court", "XX人民法院")
        doc.add_paragraph(court)

        doc.add_paragraph()
        doc.add_paragraph()

        # 签名
        p = doc.add_paragraph()
        p.add_run("原告：______________")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 日期
        from datetime import datetime
        p = doc.add_paragraph()
        p.add_run(datetime.now().strftime("%Y年%m月%d日"))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(output_path)
        return output_path

    def generate_evidence_list(self, case, output_path: str):
        """生成证据目录"""
        doc = Document()

        # 标题
        title = doc.add_paragraph("证据目录")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True

        doc.add_paragraph()

        # 案件信息
        p = doc.add_paragraph()
        p.add_run("案号：").bold = True
        p.add_run("（待分配）")

        # 获取原告和被告信息
        if hasattr(case, 'company_info'):
            plaintiff = case.company_info.target_company or ""
        else:
            plaintiff = case.get("plaintiff", "")

        if hasattr(case, 'defendants'):
            defendant = "、".join([d.def_name for d in case.defendants if d.def_name])
        else:
            defendant = case.get("defendant", "")

        p = doc.add_paragraph()
        p.add_run("原告：").bold = True
        p.add_run(plaintiff)

        p = doc.add_paragraph()
        p.add_run("被告：").bold = True
        p.add_run(defendant)

        doc.add_paragraph()

        # 证据表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # 表头
        header_cells = table.rows[0].cells
        headers = ["序号", "证据名称", "证明目的", "页码"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # 证据列表
        if hasattr(case, 'evidence'):
            evidence_list = case.evidence or ["证据1", "证据2", "证据3"]
        else:
            evidence_list = case.get("evidence", ["证据1", "证据2", "证据3"])

        for i, evidence in enumerate(evidence_list, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = evidence
            row[2].text = ""
            row[3].text = ""

        doc.save(output_path)
        return output_path

    def generate_application(self, case, content: str, output_path: str):
        """生成申请书"""
        doc = Document()

        # 标题
        title = doc.add_paragraph("申请书")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True

        doc.add_paragraph()

        # 申请人信息
        p = doc.add_paragraph()
        p.add_run("申请人：").bold = True
        if hasattr(case, 'company_info'):
            p.add_run(case.company_info.target_company or "")
        else:
            p.add_run(case.get("plaintiff", ""))

        p = doc.add_paragraph()
        p.add_run("被申请人：").bold = True
        if hasattr(case, 'defendants'):
            p.add_run("、".join([d.def_name for d in case.defendants if d.def_name]))
        else:
            p.add_run(case.get("defendant", ""))

        doc.add_paragraph()

        # 申请事项
        p = doc.add_paragraph()
        p.add_run("申请事项：").bold = True
        doc.add_paragraph("请求法院依法对本案进行审理")

        doc.add_paragraph()

        # 事实与理由
        p = doc.add_paragraph()
        p.add_run("事实与理由：").bold = True

        if hasattr(case, 'debt_info'):
            facts = f"申请人申请对被申请人进行财产保全，保全金额为{case.debt_info.guarantee_amount or ''}元。"
        else:
            facts = case.get("facts", content)

        p = doc.add_paragraph(facts if facts else "待填写")

        doc.add_paragraph()
        doc.add_paragraph()

        # 此致
        p = doc.add_paragraph()
        p.add_run("此致").bold = True
        if hasattr(case, 'case_info'):
            doc.add_paragraph(case.case_info.court_name or "XX人民法院")
        else:
            doc.add_paragraph(case.get("court", "XX人民法院"))

        doc.add_paragraph()

        # 签名
        p = doc.add_paragraph()
        p.add_run("申请人：______________")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        from datetime import datetime
        p = doc.add_paragraph()
        p.add_run(datetime.now().strftime("%Y年%m月%d日"))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(output_path)
        return output_path

    def generate_from_template(self, template_path: str, variables: Dict, output_path: str):
        """从模板生成文档，替换变量

        Args:
            template_path: 模板文件路径
            variables: 变量字典，key为变量名（不含花括号），value为替换值
            output_path: 输出文件路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        doc = Document(template_path)

        # 替换段落中的变量
        for para in doc.paragraphs:
            self._replace_paragraph_variables(para, variables)

        # 替换表格中的变量
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_paragraph_variables(para, variables)

        doc.save(output_path)
        return output_path

    def _replace_paragraph_variables(self, para, variables: Dict):
        """替换单个段落中的变量"""
        full_text = para.text

        # 检查是否包含变量
        if "{" not in full_text or "}" not in full_text:
            return

        # 构建新文本
        new_text = full_text
        for key, value in variables.items():
            # 标准格式: {variable}
            placeholder = "{" + key + "}"
            # 带空格格式: { variable } (模板中可能存在的格式问题)
            placeholder_with_space = "{ " + key + "}"

            # 替换值
            replace_value = str(value) if value else f"【{key}】"

            # 替换两种格式
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, replace_value)
            if placeholder_with_space in new_text:
                new_text = new_text.replace(placeholder_with_space, replace_value)

        # 如果文本有变化，替换段落内容
        if new_text != full_text:
            # 保存原格式
            if para.runs:
                # 获取第一个run的格式作为模板
                first_run = para.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold

                # 清空段落
                para.clear()

                # 添加新文本，保留格式
                run = para.add_run(new_text)
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = font_size
                if font_bold is not None:
                    run.font.bold = font_bold
            else:
                # 没有run，直接替换文本
                para.text = new_text

    def batch_generate_from_template(self, template_path: str, cases: list, output_dir: str) -> list:
        """批量从模板生成文档

        Args:
            template_path: 模板文件路径
            cases: 案件列表（CaseData对象）
            output_dir: 输出目录

        Returns:
            生成的文件路径列表
        """
        os.makedirs(output_dir, exist_ok=True)
        output_files = []

        for case in cases:
            # 获取模板变量
            variables = case.to_template_vars()

            # 生成文件名
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{case.company_info.target_company or case.id}_{timestamp}.docx"
            output_path = os.path.join(output_dir, output_filename)

            # 生成文档
            self.generate_from_template(template_path, variables, output_path)
            output_files.append(output_path)

        return output_files

    def extract_template_variables(self, template_path: str) -> List[str]:
        """从模板文件中提取所有变量名

        Args:
            template_path: 模板文件路径

        Returns:
            变量名列表（去重，不含花括号）
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        doc = Document(template_path)
        variables = set()

        # 从段落中提取
        for para in doc.paragraphs:
            found = re.findall(r'\{([^}]+)\}', para.text)
            for v in found:
                variables.add(v.strip())

        # 从表格中提取
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found = re.findall(r'\{([^}]+)\}', para.text)
                        for v in found:
                            variables.add(v.strip())

        return sorted(variables)

    def generate_batch_to_single_docx(self, template_path: str, data_rows: List[Dict], output_path: str):
        """批量生成文书到一个DOCX文件，每条数据单独一页

        Args:
            template_path: 模板文件路径
            data_rows: 数据行列表，每行是一个变量字典
            output_path: 输出文件路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        master_doc = Document()

        for idx, variables in enumerate(data_rows):
            # 从模板生成临时文档
            temp_doc = Document(template_path)

            # 替换临时文档中的变量
            self._replace_all_in_doc(temp_doc, variables)

            # 将临时文档的所有元素复制到主文档
            self._copy_doc_elements(temp_doc, master_doc)

            # 除最后一条外，添加分页符
            if idx < len(data_rows) - 1:
                master_doc.add_page_break()

        master_doc.save(output_path)
        return output_path

    def _replace_all_in_doc(self, doc, variables: Dict):
        """替换文档中所有变量（段落+表格）"""
        for para in doc.paragraphs:
            self._replace_paragraph_variables(para, variables)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_paragraph_variables(para, variables)

    def _copy_doc_elements(self, source_doc, target_doc):
        """将源文档的所有元素复制到目标文档（保留格式）"""
        # 复制每个段落
        for para in source_doc.paragraphs:
            new_para = target_doc.add_paragraph()
            self._copy_paragraph(para, new_para)

        # 复制表格
        for table in source_doc.tables:
            self._copy_table(table, target_doc)

        # 复制节属性（页面设置等）
        self._copy_section_properties(source_doc, target_doc)

    def _copy_paragraph(self, source_para, target_para):
        """复制段落内容和格式"""
        # 复制段落格式
        target_para.alignment = source_para.alignment
        if source_para.style:
            try:
                target_para.style = source_para.style
            except Exception:
                pass

        # 复制段落级属性
        self._copy_paragraph_properties(source_para, target_para)

        # 复制每个run
        for run in source_para.runs:
            new_run = target_para.add_run(run.text)
            self._copy_run_format(run, new_run)

        # 如果源段落有runs，清空默认空run
        if source_para.runs and target_para.runs:
            # 移除初始的空run（如果存在）
            pass

    def _copy_paragraph_properties(self, source_para, target_para):
        """复制段落属性（缩进、间距等）"""
        pPr = source_para._element.find(qn('w:pPr'))
        if pPr is not None:
            target_pPr = target_para._element.get_or_add_pPr()
            for child in list(pPr):
                tag = child.tag
                # 跳过run级别属性
                if tag in (qn('w:rPr'), qn('w:rPrChange')):
                    continue
                # 移除目标中已有的同名元素
                existing = target_pPr.findall(tag)
                for ex in existing:
                    target_pPr.remove(ex)
                target_pPr.append(copy.deepcopy(child))

    def _copy_run_format(self, source_run, target_run):
        """复制run的格式"""
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.bold is not None:
            target_run.font.bold = source_run.font.bold
        if source_run.font.italic is not None:
            target_run.font.italic = source_run.font.italic
        if source_run.font.underline is not None:
            target_run.font.underline = source_run.font.underline
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
        # 复制中文字体
        try:
            rPr_source = source_run._element.find(qn('w:rPr'))
            if rPr_source is not None:
                rPr_target = target_run._element.get_or_add_rPr()
                for child in list(rPr_source):
                    tag = child.tag
                    # 跳过已通过python-docx API设置的基本属性
                    if tag in (qn('w:rFonts'), qn('w:b'), qn('w:i'), qn('w:u'), qn('w:sz'), qn('w:color')):
                        continue
                    existing = rPr_target.findall(tag)
                    for ex in existing:
                        rPr_target.remove(ex)
                    rPr_target.append(copy.deepcopy(child))
        except Exception:
            pass

    def _copy_table(self, source_table, target_doc):
        """复制表格到目标文档"""
        # 创建新表格
        target_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))

        # 复制表格样式
        if source_table.style:
            try:
                target_table.style = source_table.style
            except Exception:
                pass

        # 复制表格对齐方式
        target_table.alignment = source_table.alignment

        # 复制每个单元格
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                target_cell = target_table.rows[i].cells[j]
                self._copy_cell(cell, target_cell)

        # 复制表格属性（宽度等）
        self._copy_table_properties(source_table, target_table)

    def _copy_cell(self, source_cell, target_cell):
        """复制单元格内容"""
        # 清除目标单元格默认段落
        for p in target_cell.paragraphs:
            p.clear()

        # 复制源单元格的每个段落
        for i, para in enumerate(source_cell.paragraphs):
            if i == 0:
                target_para = target_cell.paragraphs[0]
                target_para.clear()
            else:
                target_para = target_cell.add_paragraph()

            self._copy_paragraph(para, target_para)

        # 复制单元格属性
        self._copy_cell_properties(source_cell, target_cell)

    def _copy_cell_properties(self, source_cell, target_cell):
        """复制单元格属性（宽度、合并等）"""
        tcPr = source_cell._element.find(qn('w:tcPr'))
        if tcPr is not None:
            target_tcPr = target_cell._element.get_or_add_tcPr()
            for child in list(tcPr):
                existing = target_tcPr.findall(child.tag)
                for ex in existing:
                    target_tcPr.remove(ex)
                target_tcPr.append(copy.deepcopy(child))

    def _copy_table_properties(self, source_table, target_table):
        """复制表格属性"""
        tblPr = source_table._element.find(qn('w:tblPr'))
        if tblPr is not None:
            target_tblPr = target_table._element.get_or_add_tblPr()
            for child in list(tblPr):
                existing = target_tblPr.findall(child.tag)
                for ex in existing:
                    target_tblPr.remove(ex)
                target_tblPr.append(copy.deepcopy(child))

    def _copy_section_properties(self, source_doc, target_doc):
        """复制节属性（页面设置如页边距、纸张大小等）"""
        source_sections = source_doc.sections
        target_sections = target_doc.sections

        if not source_sections:
            return

        source_sect = source_sections[0]
        if target_sections:
            target_sect = target_sections[0]
        else:
            return

        # 复制页面设置
        if source_sect.page_width:
            target_sect.page_width = source_sect.page_width
        if source_sect.page_height:
            target_sect.page_height = source_sect.page_height
        if source_sect.top_margin:
            target_sect.top_margin = source_sect.top_margin
        if source_sect.bottom_margin:
            target_sect.bottom_margin = source_sect.bottom_margin
        if source_sect.left_margin:
            target_sect.left_margin = source_sect.left_margin
        if source_sect.right_margin:
            target_sect.right_margin = source_sect.right_margin
