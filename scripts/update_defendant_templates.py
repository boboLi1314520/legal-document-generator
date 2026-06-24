"""
修改4个模板 .docx 文件：将硬编码的2个被告段落替换为动态组合占位符
"""
from docx import Document
from docx.oxml.ns import qn
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "诉讼文书模板")

# 模板配置：文件名 -> (def1占位符替换文本, 是否需要前缀)
TEMPLATES = [
    ("1-民事起诉状-清算责任纠纷模板.docx", "{defendants_text}", True),
    ("1-民事起诉状-损害公司债权人利益责任纠纷模板.docx", "{defendants_text}", True),
    ("3-保函模板.docx", "{guarantee_defendants_text}", False),
    ("4-保全申请书模板.docx", "{respondents_text}", True),
]


def find_paragraph_index(paragraphs, marker):
    """查找包含指定标记的第一个段落索引"""
    for i, p in enumerate(paragraphs):
        if marker in p.text:
            return i
    return None


def main():
    for filename, new_placeholder, _ in TEMPLATES:
        path = os.path.join(TEMPLATE_DIR, filename)
        if not os.path.exists(path):
            print(f"文件不存在，跳过: {path}")
            continue

        print(f"处理: {filename}")
        doc = Document(path)

        # 找 def1 段落和 def2 段落
        def1_idx = find_paragraph_index(doc.paragraphs, "{def1_name}")
        def2_idx = find_paragraph_index(doc.paragraphs, "{def2_name}")

        if def1_idx is None:
            print(f"  警告: 未找到含 {{def1_name}} 的段落")
            continue
        if def2_idx is None:
            print(f"  警告: 未找到含 {{def2_name}} 的段落")
            continue

        # 1) 替换 def1 段落文本为组合占位符
        def1_para = doc.paragraphs[def1_idx]
        # 清空 def1 段落并设置新文本，保留格式
        if def1_para.runs:
            font_name = def1_para.runs[0].font.name
            font_size = def1_para.runs[0].font.size
            font_bold = def1_para.runs[0].font.bold
        else:
            font_name = None
            font_size = None
            font_bold = None

        def1_para.clear()
        run = def1_para.add_run(new_placeholder)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold

        print(f"  P{def1_idx}: 替换为 {new_placeholder}")

        # 2) 删除 def2 段落
        # 获取段落的 XML 元素并从 body 中移除
        def2_element = doc.paragraphs[def2_idx]._element
        def2_element.getparent().remove(def2_element)
        print(f"  P{def2_idx}: 已删除")

        # 保存
        doc.save(path)
        print(f"  已保存: {filename}")
        print()


if __name__ == "__main__":
    main()
